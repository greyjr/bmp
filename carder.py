#Zaporizhstal

from docx import Document
from docx.shared import Cm
import xlrd


def value_valid(ws, row, value:list):					#change float type to integer
	value = ws.cell_value(row + value[0], value[1])		
	if type(value) == float:
		value = str(int(value))
	return value


def find_replace(doc_obj, old_text, new_text, height_photo):
	for paragraph in doc_obj.paragraphs:
		if old_text in paragraph.text:
			if old_text in ['_038_photo_1_', '_039_photo_2_', '_040_photo_3_', '_041_photo_4_', '_042_photo_5_', '_043_agregat_photo_']:
				paragraph.text = ''
				paragraph.runs[0].add_picture('./mini_2/{}.jpg'.format(new_text.zfill(6)), height = Cm(height_photo))
				continue
			font_common = paragraph.runs[0].font
			text = paragraph.text
			text = text.replace(old_text, new_text)
			paragraph.text = text
			for run in range(len(paragraph.runs)):
				paragraph.runs[run].font.size = font_common.size
				paragraph.runs[run].font.name = font_common.name	
	for table in doc_obj.tables:
		for row in table.rows:
			for cell in row.cells:
				find_replace(cell, old_text, new_text, height_photo)		


def methods_fill(context):
	big_voltage = ['6кВ', '6 кВ', '0,69кВ']
	methods = ['_012_method_1_', '_017_method_2_', '_022_method_3_', '_027_method_4_', '_032_method_5_']
	values = ['_009_value_1_', '_014_value_2_', '_019_value_3_', '_024_value_4_', '_029_value_5_']
	energies = ['_008_energy_type_1_', '_013_energy_type_2_', '_018_energy_type_3_', '_023_energy_type_4_', '_028_energy_type_5_']
	methods_text = [	'Визуальный осмотр, проверить отсутствие напряжения (тестер) и пробное включение механизма после блокировки органами управления в присутствии производителя работ',
				'Визуальный осмотр и пробное включение механизма после блокировки органами управления в присутствии производителя работ',
				'После закрытия стравить избыточное давление (при наличии дренажа оставить его в открытом состоянии). Визуальный осмотр, попытка открытия после блокировки в присутствии производителя работ',
				]
	for i in range(5):
		if context[energies[i]] != 'электричество':
			context[methods[i]] = methods_text[2]
		elif context[values[i]] in big_voltage:
			context[methods[i]] = methods_text[1]
		else:
			context[methods[i]] = methods_text[0]
			

def write_word(context:dict, filename:str, size_card:int):											#
	doc = Document(filename)
	methods_fill(context)
	height_photo = 5.55 if size_card>1 else 10
	for key in context:
		find_replace(doc, key, context[key], height_photo)
	

	token = ['_010_blockirator_1_', '_015_blockirator_2_', '_020_blockirator_3_', '_025_blockirator_4_', '_030_blockirator_5_']
	block_list_keys = []
	for i in range(size_card):
		block_list_keys.append(token[i])
	block_list = list(set([context[i] for i in block_list_keys]))
	
	name = "_".join(sorted(filter(None,block_list)))
	new_p = doc.add_paragraph()
	new_r = new_p.add_run()
	new_r.add_picture('./block/{}{}.jpg'.format(name,(size_card+3)//3), width = Cm(27.5))



	card_name = context['_000_card_number_'] + ' ' + context['_006_agregat_name_']
	doc.save('./cards/{}.docx'.format(card_name))
	print('Card N {} saved...'.format(card_name))
	

def main():
	wb = xlrd.open_workbook('общая БМП.xlsx')
	ws = wb.sheet_by_name('ДЦ ЗС 1')
	begin = 4
	end = 40

	card_context_list = {
		'_000_card_number_':[0,0],
		'_001_point_1_':[0,4],
		'_002_point_2_':[1,4],
		'_003_point_3_':[2,4],
		'_004_point_4_':[3,4],
		'_005_point_5_':[4,4],
		'_006_agregat_name_':[0,1],
		'_007_section_':[0,12],
		'_008_energy_type_1_':[0,2],
		'_009_value_1_':[0,6],
		'_010_blockirator_1_':[0,9],
		'_011_lock_1_':[0,10],
		'_012_method_1_':[0,0],#METHOD !!!!!!!!!
		'_013_energy_type_2_':[1,2],
		'_014_value_2_':[1,6],
		'_015_blockirator_2_':[1,9],
		'_016_lock_2_':[1,10],
		'_017_method_2_':[0,0],
		'_018_energy_type_3_':[2,2],
		'_019_value_3_':[2,6],
		'_020_blockirator_3_':[2,9],
		'_021_lock_3_':[2,10],
		'_022_method_3_':[0,0],
		'_023_energy_type_4_':[3,2],
		'_024_value_4_':[3,6],
		'_025_blockirator_4_':[3,9],
		'_026_lock_4_':[3,10],
		'_027_method_4_':[0,0],
		'_028_energy_type_5_':[4,2],
		'_029_value_5_':[4,6],
		'_030_blockirator_5_':[4,9],
		'_031_lock_5_':[4,10],
		'_032_method_5_':[0,0],
		'_033_location_1_':[0,5],
		'_034_location_2_':[1,5],
		'_035_location_3_':[2,5],
		'_036_location_4_':[3,5],
		'_037_location_5_':[4,5],
		'_038_photo_1_':[0,8],
		'_039_photo_2_':[1,8],
		'_040_photo_3_':[2,8],
		'_041_photo_4_':[3,8],
		'_042_photo_5_':[4,8],
		'_043_agregat_photo_':[0,7],
		}
	row = begin
	break_flag = False
	while row<=end:
		context = {record:value_valid(ws,row,value) for record,value in card_context_list.items()}	#read excel and create dict CONTEXT
		context['_000_card_number_'] = context['_000_card_number_'].zfill(3)						#fill number by 000    N3 --> N003

		size_card = 1																				#min size card
		while not ws.cell_value(row+size_card,0):													#calculate real card size
			size_card+=1
			if size_card > 16:																		#catch error (16 - max)
				break_flag =True
				break
		if break_flag:
			break
		filename = 'blank_{}.docx'.format(size_card)												#name new docx-file
		if size_card>5:																				#I haven't blanks for 6 and more points
			print('\nCard N {} to big! Skip.\n'.format(context['_000_card_number_']))				#message to console if skip
			continue
			print('Card N {} are big. Skip\n'.format(context['_000_card_number_']))
			continue
		if context['_006_agregat_name_'][0] == '*':													#skip marked card
			print('Card N {} marked as not ready. Skip\n'.format(context['_000_card_number_']))
			continue
		
		write_word(context, filename, size_card)													#call write-procedure
		row = row + size_card																		#next row

	print('***********finished***********')


if __name__ == '__main__':
	main()

